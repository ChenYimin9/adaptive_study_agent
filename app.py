# Import various modules
import streamlit as st
import random  
import pandas as pd
import json
import uuid
from datetime import datetime, timedelta
from streamlit_autorefresh import st_autorefresh 
from core.backend import (
    MORANDI_COLORS, extract_text_from_file, MockUserManager, MockLearningEngine,
    MockAssessmentManager, DeepSeekAIAgent, MockAssistanceTracker, 
    MockPDFGenerator, MockLearningAnalytics as LearningAnalytics
)
from core.user_manager import user_manager

FIXED_MOTIVATIONAL_MESSAGES = [
    {
        "message": "Perseverance leads to victory! \n You have accumulated a considerable amount of time for study. \n Keep up this pace and the goal is just ahead.",
        "quote": "The secret of success lies in consistently adhering to one's goals. - Disraeli"
    },
    {
        "message": "Every learning experience is an accumulation of progress. \n Today's efforts will become the confidence for tomorrow.",
        "quote": "Learning is a continuous process with no end. -- Einstein"
    },
    {
        "message": "Don't underestimate the little progress made every day. \n With the accumulation of time, it will eventually lead to a qualitative leap. \n Keep it up!",
        "quote": "A journey of a thousand miles begins with a single step. -- Laozi"
    },
    {
        "message": "You are cultivating excellent study habits, \n which are more valuable than any knowledge.",
        "quote": "Habit is a tenacious and powerful force that can dominate one's life. -- Bacon"
    },
    {
        "message": "Don't back down when you encounter difficulties. This is precisely the opportunity for growth. \n Looking back on what you have learned, \n you will find that you have long surpassed your past self.",
        "quote": "Difficulties are like springs. \n The stronger you are, the weaker they become; \n the weaker you are, the stronger they become. -- Anonymous"
    },
    {
        "message": "Time is the most fair judge. \n Every minute you invest in learning \n will bloom into unexpected results in the future.",
        "quote": "Lost time is never found again. -- Franklin"
    },
    {
        "message": "Mistakes are not failures, but signposts. \n They tell you where to focus next \n and make your learning more effective.",
        "quote": "I have not failed. I've just found 10,000 ways that won't work. -- Edison"
    },
    {
        "message": "Even on days when motivation is low, showing up is a victory. \n Consistency beats intensity in the long run.",
        "quote": "Success is the sum of small efforts, repeated day in and day out. -- Rohn"
    },
    {
        "message": "Your curiosity is the best teacher. \n Keep asking 'why' and exploring 'how' \n and you'll unlock more knowledge than you imagine.",
        "quote": "The important thing is not to stop questioning. -- Einstein"
    },
    {
        "message": "Celebrate every small win. \n They are the building blocks of big achievements \n and will fuel your journey forward.",
        "quote": "Enjoy the little things, for one day you may look back and realize they were the big things. -- Stevenson"
    }
]

# Remove_topic_questions
def mock_remove_topic_questions(self, path_id, topic_name, user_id):
    return True
MockLearningEngine.remove_topic_questions = mock_remove_topic_questions

# Initialize the core module (globally unique)
learning_engine = MockLearningEngine()
assessment_manager = MockAssessmentManager()
PDFGenerator = MockPDFGenerator()

# Page configuration
st.set_page_config(
    page_title="Adaptive Study Companion",
    layout="wide",
    initial_sidebar_state="expanded",
    page_icon="📚"
)

# Page style
st.markdown(f"""
<style>
    :root {{
        --primary: {MORANDI_COLORS['primary']};
        --secondary: {MORANDI_COLORS['secondary']};
        --accent: {MORANDI_COLORS['accent']};
        --light: {MORANDI_COLORS['light']};
        --success: {MORANDI_COLORS['success']};
        --warning: {MORANDI_COLORS['warning']};
        --text: {MORANDI_COLORS['text']};
        --card_bg: {MORANDI_COLORS['card_bg']};
        --bg: {MORANDI_COLORS['bg']};
        --module-spacing: 25px;
    }}
    
    /* Basic font: Enlarge by 10% to ensure clear text */
    body {{ background-color: var(--bg); color: var(--text); font-size: 1.65rem; margin: 0; padding: 0; line-height: 1.7; }}
    .main-header {{ font-size: 3rem; color: var(--accent); margin-bottom: 1.8rem; font-weight: 700; line-height: 1.2; }}
    
    /* Title hierarchy: Fine-tune the desktop size to maintain a sense of hierarchy */
    h1, h2, h3, h4, h5, h6 {{
        color: var(--accent);
        margin-top: 1.8rem;
        margin-bottom: 1rem;
    }}
    h1 {{ font-size: 2.8rem; }} h2 {{ font-size: 2.6rem; }} h3 {{ font-size: 2.4rem; }} h4 {{ font-size: 2.2rem; }} h5 {{ font-size: 2rem; }}
    
    /* Button font: On mobile devices, enlarge from 1.4rem to 1.6rem to avoid being too small */
    .stButton > button, [data-baseweb="button"], .stFormSubmitButton > button {{
        background-color: var(--primary) !important; color: white !important; border-radius: 10px !important;
        padding: 1rem 2rem !important; font-size: 1.6rem !important; border: none !important;
        margin: 0.6rem !important; box-shadow: 0 3px 10px rgba(0,0,0,0.12) !important;
        min-height: 55px !important; white-space: nowrap !important;
    }}
    
    /* Input box: Adjust the label from 2rem to 1.8rem (not jarring), and keep the input content at 2rem (clear) */
    .stTextInput > label, .stTextArea > label, .stSelectbox > label, .stSlider > label {{
        font-size: 1.8rem !important; font-weight: 500; margin-bottom: 0.8rem; display: block;
    }}
    .stTextInput > div > div > input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {{
        font-size: 2rem !important; padding: 1rem !important; border-radius: 10px;
        border: 1px solid #ddd; height: auto; min-height: 55px !important;
    }}
    
    /* System text: Enlarge from 1.4rem to 1.5rem to ensure the clarity of the auxiliary text */
    .css-1d391kg {{ padding-top: 2rem; border-top: 2px solid #0078D4; }}
    .css-1v3fvcr, .css-1dimb5e, .css-10trblm {{ font-size: 1.5rem !important; line-height: 1.7; }}
    
    /* Folding panel: On mobile devices, zoom in from 1.4rem to 1.6rem to avoid unclear view when clicking */
    .stExpander > button {{ font-size: 2.2rem !important; font-weight: 500; padding: 1rem; }}
    .stExpander > div {{ padding: 1rem 1.2rem; }}
    
    /* TAB: On mobile devices, zoom in from 1.3rem to 1.4rem to ensure clear switching options */
    .stTabs [data-baseweb="tab-list"] button {{ font-size: 1.5rem !important; padding: 1rem 1.8rem; }}
    
    /* Data indicators: The label maintains 1.4rem (not stealing the spotlight), and the value is 2rem (highlighting key points) */
    .stMetric {{ padding: 1.2rem; }}
    .stMetric label {{ font-size: 1.4rem !important; }}
    .stMetric value {{ font-size: 2rem !important; }}
    
    .stForm {{ padding: 1.5rem; }}
    
    /* Chart text: Enlarge as a whole to ensure clear chart annotations */
    .plotly-graph-div {{ font-size: 1.5rem !important; }}
    .js-plotly-plot .plotly .legend {{ font-size: 1.7rem !important; }}
    .js-plotly-plot .plotly .axis-title {{ font-size: 1.9rem !important; }}
    .js-plotly-plot .plotly .xtick, .js-plotly-plot .plotly .ytick {{ font-size: 1.5rem !important; }}

    /* Mobile responsive: The core is that all fonts should not be smaller than 1.4rem to avoid being too small */
    @media (max-width: 768px) {{
        .main-header {{ font-size: 2.6rem; }}
        h1 {{ font-size: 2.4rem; }} h2 {{ font-size: 2.2rem; }} h3 {{ font-size: 2rem; }}
        
        .stButton > button, [data-baseweb="button"], .stFormSubmitButton > button {{
            padding: 0.8rem 1.5rem !important; font-size: 1.6rem !important; min-height: 50px !important;
        }}
        
        .stTextInput > div > div > input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {{
            padding: 0.8rem !important; min-height: 50px !important;
        }}
        
        .stTextInput > label, .stTextArea > label, .stSelectbox > label, .stSlider > label {{ font-size: 1.6rem !important; }}
        .stExpander > button {{ font-size: 1.6rem !important; padding: 0.8rem; }}
        .stExpander > div {{ padding: 0.8rem 1rem; }}
        .stTabs [data-baseweb="tab-list"] button {{ font-size: 1.4rem !important; padding: 0.8rem 1.2rem; }}
        .stMetric {{ padding: 1rem; }}
        .stMetric label {{ font-size: 1.4rem !important; }}
        .stMetric value {{ font-size: 1.8rem !important; }}
        .stForm {{ padding: 1rem; }}
        .plotly-graph-div {{ font-size: 1.5rem !important; }}
        .js-plotly-plot .plotly .legend {{ font-size: 1.5rem !important; }}
        .js-plotly-plot .plotly .axis-title {{ font-size: 1.6rem !important; }}
        .js-plotly-plot .plotly .xtick, .js-plotly-plot .plotly .ytick {{ font-size: 1.4rem !important; }}
    }}
    
    /* Function Style: Ensure that the font of custom components is not too small */
    .delete-btn {{ background-color: #ff4b4b !important; color: white !important; padding: 0.5rem 1rem !important; font-size: 1.5rem !important; min-height: auto !important; margin: 0 !important; }}
    
    .path-card {{
        background-color: var(--card_bg); border-radius: 10px; padding: 1.5rem; margin-bottom: 1.5rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08); border-left: 4px solid var(--primary);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }}
    .path-card:hover {{ transform: translateY(-3px); box-shadow: 0 6px 16px rgba(0,0,0,0.12); }}
    .path-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 1rem; }}
    .path-title {{ font-size: 2rem; font-weight: 600; color: var(--primary); }} /* "Path title enlarged */
    .path-progress {{ font-size: 1.6rem; font-weight: 500; color: var(--accent); }} /* Progress text enlargement */
    .path-details {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 1rem; margin-bottom: 1.5rem; }}
    .path-detail-item {{ font-size: 1.5rem; }} /* Path details zoom in */
    .path-actions {{ display: flex; gap: 1rem; margin-top: 1rem; }}
    .delete-path-btn {{ background-color: #ff4b4b !important; }}
    
    .exercise-card {{ background-color: var(--light); border-radius: 8px; padding: 1.2rem; margin: 1rem 0; border-left: 4px solid var(--accent); }}
    .exercise-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem; }}
    .exercise-title {{ font-size: 1.7rem; font-weight: 600; color: var(--primary); }} /* The title of the exercise is enlarged */
    .exercise-difficulty {{ font-size: 1.5rem; font-weight: 500; padding: 0.3rem 0.8rem; border-radius: 12px; background-color: var(--warning); color: white; }} /* 难度标签放大 */
    .exercise-content {{ margin: 0.8rem 0; }}
    .exercise-answer {{ margin-top: 1rem; padding: 0.8rem; background-color: rgba(98, 100, 167, 0.1); border-radius: 6px; }}
    
    .feedback-correct {{ background-color: #d4edda; color: #155724; padding: 0.8rem; border-radius: 6px; margin-top: 1rem; border-left: 4px solid #28a745; }}
    .feedback-incorrect {{ background-color: #f8d7da; color: #721c24; padding: 0.8rem; border-radius: 6px; margin-top: 1rem; border-left: 4px solid #dc3545; }}
    
    .message-container {{ display: flex; flex-direction: column; margin-bottom: 1.5rem; border: 1px solid #e0e0e0; border-radius: 8px; padding: 1rem; background-color: white; }}
    .message-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem; }}
    .message-role {{ font-weight: bold; color: var(--primary); }}
    .message-timestamp {{ font-size: 1.2rem; color: #777; }} /* The timestamp is enlarged from 0.9rem to 1.2rem to avoid being too small */
    .message-content {{ margin-top: 0.5rem; line-height: 1.5; }}
    .message-actions {{ display: flex; justify-content: flex-end; margin-top: 0.5rem; }}
    
    .assessment-container {{ border: 1px solid #e0e0e0; border-radius: 10px; padding: 1.5rem; margin: 1.5rem 0; background-color: #f9f9f9; }}
    .assessment-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 1.5rem; }}
    .assessment-title {{ font-size: 2rem; font-weight: 600; color: var(--primary); }} /* Enlarge the evaluation title */
    .assessment-score {{ font-size: 2rem; font-weight: 700; color: var(--accent); }} /* "Score amplification" */
    
    .question-container {{ margin-bottom: 2rem; padding: 1.2rem; background-color: white; border-radius: 8px; box-shadow: 0 2px 6px rgba(0,0,0,0.1); }}
    .question-header {{ display: flex; justify-content: space-between; margin-bottom: 1rem; }}
    .question-number {{ font-size: 1.6rem; font-weight: 600; color: var(--primary); }} /* Enlarge the question number */
    .question-difficulty {{ font-size: 1.5rem; font-weight: 500; padding: 0.3rem 0.8rem; border-radius: 12px; background-color: var(--warning); color: white; }} /* 题目难度放大 */
    .question-content {{ font-size: 1.7rem; margin-bottom: 1.2rem; line-height: 1.6; }} /* Enlarge the content of the question */
    .feedback-section {{ padding: 1rem; border-radius: 8px; margin-top: 1rem; }}
    .feedback-title {{ font-weight: 600; margin-bottom: 0.5rem; }}
    .answer-section {{ padding: 1rem; background-color: #f0f8ff; border-radius: 8px; margin-top: 1rem; }}
    
    .viewed-btn {{ background-color: #cccccc !important; color: #666666 !important; pointer-events: none; }}
</style>
""", unsafe_allow_html=True)

# Session State Initialization
def init_session_state():
    session_vars = {
        'user': None, 'current_view': 'dashboard', 'current_assessment': None,
        'active_path': None, 'loading': False, 'notifications': [], 'last_action': None,
        'last_login': None, 'completed_paths': [], 'selected_path_id': None,
        'uploaded_materials': {}, 'assessment_state': {}, 'assistance_requests': [],
        'chat_history': [], 'api_key': "", 'assessment_results': [], 'path_actions': {},
        'show_progress_update': {}, 'ai_generated_path': False, 'delete_message_id': None,
        'expanded_exercises': {}, 'viewed_resources': {}, 'exercise_feedback': {},
        'exercise_answers': {}, 'exercise_submitted': {}, 'topic_assessments': {},
        'resource_views': {}, 'saved_assessments': {}, 'achievements': [],
        'path_resource_counts': {}, 'path_topic_counts': {}, 'topic_progress': {},
        'achievement_path_ids': set(), 'assessment_generating': {}, 'show_assessment': {},
        'topic_assessment_scores': {},
        'timer_initialized': False,  
        'last_timer_update': None     
    }
    for var, val in session_vars.items():
        if var not in st.session_state:
            st.session_state[var] = val

init_session_state()

# Initialize the AI agent
def init_ai_agent():
    if 'ai_agent' not in st.session_state:
        st.session_state.ai_agent = DeepSeekAIAgent(api_key=st.session_state.api_key)
init_ai_agent()

# Calculate the progress of the learning path
def calculate_path_progress(path_id, topics, path):
    user = st.session_state.user
    if not user: return 0.0
    
    total_resources = viewed_resources = 0
    total_topics = len(topics)
    completed_topics = 0
    topic_progress_list = []
    
    for i, topic in enumerate(topics):
        resources = topic.get('resources', [])
        topic_resources = len(resources)
        total_resources += topic_resources
        
        viewed_in_topic = 0
        for res in resources:
            if learning_engine.check_resource_viewed(user['id'], path_id, topic['name'], res['title']):
                viewed_in_topic += 1
                viewed_resources += 1
        
        resource_progress = viewed_in_topic / topic_resources if topic_resources > 0 else 1.0
        topic_key = f"{path_id}_topic_{i}_{topic['name']}"
        state_content = learning_engine.get_assessments_by_topic(user['id'], path['subject'], topic['name'])
        
        if state_content and state_content.get('content'):
            scores = json.loads(state_content['content'])['scores']
            valid_scores = [s for s in scores if isinstance(s, (int, float))]
            score = (sum(valid_scores)/len(valid_scores)*100) if valid_scores else 0.0
            assessment_passed = score >= 80
            topic_completed = assessment_passed and resource_progress >= 0.8
            topic_progress = 1.0 if topic_completed else resource_progress
            topic_progress_list.append(topic_progress)
            if topic_completed: completed_topics += 1
    
    overall_progress = completed_topics / total_topics if total_topics > 0 else 0
    
    st.session_state.path_resource_counts[path_id] = {'total': total_resources, 'viewed': viewed_resources}
    st.session_state.path_topic_counts[path_id] = {'total': total_topics, 'completed': completed_topics}
    st.session_state.topic_progress[path_id] = topic_progress_list
    
    return round(min(1.0, overall_progress), 4)

# -------------------------- Real-time timing is connected to the back end --------------------------
def get_default_path_topic(user_id):
    """Obtain the default path_id and topic_name"""
    paths = learning_engine.get_learning_paths(user_id)
    if not paths:
        # Create a default path when there is no path
        default_path_id = learning_engine.create_learning_path(
            user_id, "Default Study", "Beginner", 30, st.session_state.ai_agent
        )[0]
        path = learning_engine.get_learning_path(default_path_id, user_id)
        content = json.loads(path['content']) if path else {"topics": [{"name": "Default Topic"}]}
        return default_path_id, content['topics'][0]['name'] if content['topics'] else "Default Topic"
    
    # When there is a path, use the first topic of the first path
    first_path = paths[0]
    content = json.loads(first_path['content'])
    first_topic = content['topics'][0]['name'] if content['topics'] else "Default Topic"
    return first_path['id'], first_topic

def init_backend_timer():
    """Initialize the backend timing"""
    if st.session_state.user and not st.session_state.timer_initialized:
        user_id = st.session_state.user['id']
        # Get the default path and topic
        path_id, topic_name = get_default_path_topic(user_id)
        # Call the backend initialization interface
        init_result = learning_engine.init_study_timer(user_id, path_id, topic_name)
        if init_result["status"] == "success":
            st.session_state.timer_initialized = True
            st.session_state.last_timer_update = datetime.now()
            st.success("Total study duration may have a delay.")
        else:
            st.warning(f"The timing initialization failed：{init_result['message']}")

def update_backend_timer():
    """Update the backend timing"""
    if not (st.session_state.user and st.session_state.timer_initialized):
        return 0.0  
    # Return 0 if not initialized
    
    user_id = st.session_state.user['id']
    current_time = datetime.now()
    # Control the invocation frequency
    if (current_time - st.session_state.last_timer_update).total_seconds() < 10:
        # Check the latest duration directly without repeated updates
        total_min = learning_engine.get_total_study_time(user_id)
        return round(total_min, 2)
    
    # 1. Clearly define the number of minutes corresponding to 10 seconds using floating-point numbers (10 seconds =0.1667 minutes)
    add_minutes = 10.0 / 60  # Use 10.0 to ensure that the floating-point number calculation results ≈0.1667
    add_minutes = round(add_minutes, 2)  # Retain two decimal places and pass it on for 0.17 minutes
    
    # 2. Get the default path and topic
    path_id, topic_name = get_default_path_topic(user_id)
    
    # 3. Call the backend update interface
    update_result = learning_engine.update_study_timer(user_id, path_id, topic_name, add_minutes)
    if update_result["status"] == "success":
        st.session_state.last_timer_update = current_time
        # Return the latest total duration (rounded to two decimal places for a more friendly display)
        return round(update_result["total_study_time"], 2)
    else:
        st.warning(f"Duration update failed：{update_result['message']}（Current transmission duration:{add_minutes}Minutes）")
        return round(learning_engine.get_total_study_time(user_id), 2)

def get_total_study_time_from_backend():
    """Obtain the total learning duration from the back end"""
    if not st.session_state.user:
        return 0
    # First, try to update, and then return the latest duration
    return update_backend_timer()

# -------------------------- Page component --------------------------
# AI Assistant interface
def show_ai_assistant():
    # Initialize the backend timing
    init_backend_timer()
    # Update and get the total duration
    total_minutes = get_total_study_time_from_backend()
    total_hours = total_minutes / 60  # Convert to hours
    
    user = st.session_state.user
    if not user:
        st.session_state.current_view = 'login'
        st.rerun()
        
    # Display the real-time total duration
    st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m", delta="There may be delays.")   
    st.subheader("🤖 AI Assistant")
    st.write("Ask any questions about your studies and get instant help")

   # Handle message deletion
    if st.session_state.delete_message_id is not None:
        st.session_state.chat_history = [
            msg for i, msg in enumerate(st.session_state.chat_history) 
            if i != st.session_state.delete_message_id
        ]
        st.session_state.delete_message_id = None
        st.rerun()
    
    # Display chat history
    if st.session_state.chat_history:
        st.write("### Conversation History")
        
        if st.button("🗑️ Delete All History", use_container_width=True):
            st.session_state.chat_history = []
            st.success("All chat history has been deleted")
            st.rerun()
        
        for i, message in enumerate(st.session_state.chat_history):
            role = "You" if message["role"] == "user" else "Assistant"
            color = MORANDI_COLORS['primary'] if message["role"] == "assistant" else MORANDI_COLORS['accent']
            
            with st.container():
                st.markdown(f"""
                <div class="message-container">
                    <div class="message-header">
                        <div class="message-role" style="color: {color};">{role}</div>
                        <div class="message-timestamp">{message['timestamp']}</div>
                    </div>
                    <div class="message-content">{message['content']}</div>
                    <div class="message-actions">
                """, unsafe_allow_html=True)
                
                if st.button("🗑️ Delete", key=f"delete_{i}"):
                    st.session_state.delete_message_id = i
                    st.rerun()
                
                st.markdown("</div></div>", unsafe_allow_html=True)
        
        st.markdown("---")
    
    # User input
    with st.form("ai_assistant_form", clear_on_submit=True):
        user_input = st.text_area("Your question:", placeholder="Ask anything about your studies...", height=150)
        submit = st.form_submit_button("Send")
        
        if submit and user_input:
            user_msg = {
                "role": "user", "content": user_input, "timestamp": datetime.now().strftime("%H:%M")
            }
            st.session_state.chat_history.append(user_msg)
            
            conversation_history = [{"role": msg["role"], "content": msg["content"]} 
                                   for msg in st.session_state.chat_history]
            
            with st.spinner("Thinking..."):
                try:
                    if not st.session_state.api_key:
                        st.error("DeepSeek API Key is required. Please set it in the dashboard.")
                        return
                    
                    response = st.session_state.ai_agent.chat(conversation_history)
                    
                    if response and isinstance(response, dict) and 'error' not in response:
                        assistant_msg = {
                            "role": "assistant", "content": response['response'], 
                            "timestamp": datetime.now().strftime("%H:%M")
                        }
                        st.session_state.chat_history.append(assistant_msg)
                        st.rerun()
                    else:
                        error_msg = response.get('error', 'Failed to get response') if isinstance(response, dict) else 'Unknown error'
                        st.error(f"Assistant error: {error_msg}")
                except Exception as e:
                    st.error(f"Error communicating with AI: {str(e)}")

# Login page
def show_login_page():
    col_banner, col_form = st.columns([1, 1])
    
    with col_banner:
        st.markdown(f"""
        <div style="background-color: {MORANDI_COLORS['primary']}; color: white; padding: 3rem;
                    border-radius: 15px 0 0 15px; height: 100%; display: flex; flex-direction: column; 
                    justify-content: center;">
            <h1 class="main-header" style="color: white;">📚 Adaptive Study Companion</h1>
            <p style="font-size: 1.5rem; line-height: 1.7; margin-bottom: 2.5rem;">
                Personalized learning paths tailored to your goals. Track progress, identify strengths, and achieve mastery at your own pace.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col_form:
        st.subheader("Log in")
        with st.form("login_form", clear_on_submit=True):
            username = st.text_input("Username", key="login_user")
            password = st.text_input("Password", type="password", key="login_pass")
            submit_login = st.form_submit_button("Login")

            if submit_login:
                with st.spinner("Authenticating..."):
                    user = user_manager.authenticate_user(username, password)
                    if user:
                        # Login successful: Initialize session status + backend timing
                        st.session_state.user = user
                        st.session_state.current_view = 'dashboard'
                        st.session_state.last_login = datetime.now()
                        st.session_state.timer_initialized = False 
                        st.success("Login successful!")
                        st.rerun()
                    else:
                        st.error("Invalid username or password")
        
        st.markdown("---")
        
        st.subheader("New user? Register")
        with st.form("register_form", clear_on_submit=True):
            col1, col2 = st.columns(2, gap="small")
            with col1:
                new_username = st.text_input("Username", key="reg_user")
                new_email = st.text_input("Email", key="reg_email")
            with col2:
                new_password = st.text_input("Password", type="password", key="reg_pass")
                full_name = st.text_input("Full name", key="reg_name")
            
            interests = st.text_area("Interests (comma-separated)", key="reg_interests", height=150)
            learning_style = st.selectbox(
                "Learning style", ["Visual", "Auditory", "Reading/Writing", "Kinesthetic"], key="reg_style"
            )
            
            submit_reg = st.form_submit_button("Register")
            
            if submit_reg:
                with st.spinner("Creating account..."):
                    user_id = user_manager.register_user(
                        new_username, new_password, new_email, full_name, interests, learning_style
                    )
                    if user_id:
                        st.success("Registration successful! Please log in.")
                    else:
                        st.error("Registration failed. Try a different username.")

# Dashboard page
def show_dashboard():
    # Initialize the backend timing
    init_backend_timer()
    # Update and get the total duration
    total_minutes = get_total_study_time_from_backend()
    total_hours = total_minutes / 60  # Convert to hours and retain one decimal place
    
    user = st.session_state.user
    if not user:
        st.session_state.current_view = 'login'
        st.rerun()
        
    st.markdown(f'<h1 class="main-header">📚 Welcome back, {user["username"]}!</h1>', unsafe_allow_html=True)
    
    # API Key Settings
    with st.expander("🔑 API Settings", expanded=False):
        new_api_key = st.text_input("DeepSeek API Key", value=st.session_state.get('api_key', ''), type="password")
        if st.button("Save API Key"):
            st.session_state.api_key = new_api_key
            st.session_state.ai_agent.api_key = new_api_key
            st.success("API Key saved successfully!")
    
   # Certificate Download
    if st.session_state.completed_paths:
        with st.expander("🏆 You have completed courses! Download certificates", expanded=True):
            for idx, path in enumerate(st.session_state.completed_paths):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"{path['subject']} - Completed on {datetime.now().strftime('%Y-%m-%d')}")
                with col2:
                    if st.button("Download Certificate", key=f"cert_btn_{path['id']}_{idx}"):
                        user_profile = user_manager.get_user_profile(user['id'])
                        cert_result = PDFGenerator.generate_certificate(
                            user['full_name'], path['subject'], datetime.now().strftime('%Y-%m-%d')
                        )
                        if cert_result["status"] == "success":
                            st.download_button(
                                "Save Certificate", cert_result["buffer"],
                                f"certificate_{path['subject'].replace(' ', '_')}_{cert_result['cert_number']}.pdf",
                                "application/pdf", key=f"download_cert_{path['id']}_{idx}"
                            )
    
    # Display the real-time total duration obtained by the backend
    analytics = learning_engine.get_learning_analytics(
        user['id'], st.session_state.completed_paths  # Remove the assessment_results of the session status
    )
    if isinstance(analytics, dict) and "status" not in analytics: # Ensure that analytics is valid data
        # Login time display
        last_login_time = st.session_state.last_login.strftime("%Y-%m-%d %H:%M") if st.session_state.last_login else "N/A"
        
        # Display the real-time total duration (obtained from the back end)
        cols = st.columns(2)
        with cols[0]: st.metric("Last Login", last_login_time)
        with cols[1]: st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m")
    

    # Randomly select an incentive message from a fixed list (cache for 1 hour to avoid frequent changes)
    if "motivation_data" not in st.session_state or (
        datetime.now() - st.session_state.get("motivation_cache_time", datetime.now())
    ).total_seconds() > 3600:
        # Randomly select one from the fixed list
        st.session_state["motivation_data"] = random.choice(FIXED_MOTIVATIONAL_MESSAGES)
        st.session_state["motivation_cache_time"] = datetime.now()  # Record cache time
    # Display fixed incentive information
    st.info(f"{st.session_state['motivation_data']['message']}\n\n*{st.session_state['motivation_data']['quote']}*")
        
    # Learning path display
    st.subheader("Your Learning Paths")
    paths = learning_engine.get_learning_paths(user['id'])
    
    if paths:
        for path in paths:
            with st.container():
                st.markdown(f"""
                <div class="path-card">
                    <div class="path-header">
                        <div class="path-title">{path['subject']}</div>
                        <div class="path-progress">{path['progress']*100:.1f}% complete</div>
                    </div>
                    <div class="path-details">
                        <div class="path-detail-item"><strong>Difficulty:</strong> {path['difficulty_level']}</div>
                        <div class="path-detail-item"><strong>Target:</strong> {path['target_completion_date']}</div>
                        <div class="path-detail-item"><strong>Created:</strong> {path['created_at'].strftime('%Y-%m-%d')}</div>
                        <div class="path-detail-item"><strong>Last Updated:</strong> {path['last_updated'].strftime('%Y-%m-%d')}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("Update Progress", key=f"update_{path['id']}"):
                        st.session_state.show_progress_update[path['id']] = True
                        st.rerun()
                with col2:
                    if st.button("Delete Path", key=f"delete_{path['id']}"):
                        learning_engine.delete_learning_path(path['id'], user['id'])
                        st.success("Learning path deleted successfully!")
                        st.rerun()
                with col3:
                    if st.button("Continue", key=f"continue_{path['id']}"):
                        st.session_state.active_path = learning_engine.get_learning_path(path['id'], user['id'])
                        st.session_state.selected_path_id = path['id']
                        st.session_state.show_assessment = {}
                        st.session_state.current_view = 'learning_path'
                        st.rerun()
                
                if st.session_state.show_progress_update.get(path['id'], False):
                    new_progress = st.slider(
                        f"Update progress for {path['subject']}", 0.0, 1.0, path['progress'], key=f"prog_{path['id']}"
                    )
                    if st.button(f"Save Progress", key=f"save_{path['id']}"):
                        learning_engine.update_learning_progress(path['id'], user['id'], new_progress)
                        st.session_state.show_progress_update[path['id']] = False
                        st.success("Progress updated!")
                        st.rerun()
    else:
        st.info("No learning paths yet. Create one to start!")
    
    # Create Path Button
    if st.button("➕ Create New Learning Path", use_container_width=True):
        st.session_state.current_view = 'create_path'
        st.rerun()
    
    # Activity Heat Map
    if isinstance(analytics, dict) and "activities" in analytics and analytics["activities"]:
       st.subheader("📊 Your Activity")
    # Extract the Plotly Figure object from the dictionary
    heatmap_result = LearningAnalytics.generate_activity_heatmap(analytics['activities'])
    if heatmap_result and heatmap_result.get("status") == "success":
        fig = heatmap_result["figure"]  # Only take the Figure object
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Failed to generate activity heatmap")
    
    # Learning suggestions
    if isinstance(analytics, dict) and "assessments" in analytics:
        analysis_data = assessment_manager.get_weakness_areas(user['id'], analytics['assessments'])
        if analysis_data and "status" not in analysis_data:
            st.subheader("🎯 Learning Recommendations")
            
            if analysis_data.get('strong_topics'):
                with st.expander("✅ Your Strengths", expanded=True):
                    st.markdown("""
                    <div class="strength-card"><div class="analysis-header"><i>✅</i> You excel in these areas:</div>
                    """, unsafe_allow_html=True)
                    for topic in analysis_data['strong_topics'][:3]:
                        st.markdown(f"""
                        <div class="topic-item">
                            <span class="topic-name">{topic['topic']}</span>
                            <span class="topic-score">{topic['avg_score']*100:.1f}%</span>
                        </div>
                        """, unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)
            
            if analysis_data.get('weak_topics'):
                with st.expander("⚠️ Areas to Improve", expanded=True):
                    st.markdown("""
                    <div class="weakness-card"><div class="analysis-header"><i>⚠️</i> Focus on these topics for improvement:</div>
                    """, unsafe_allow_html=True)
                    for topic in analysis_data['weak_topics'][:3]:
                        st.markdown(f"""
                        <div class="topic-item">
                            <span class="topic-name">{topic['topic']}</span>
                            <span class="topic-score">{topic['avg_score']*100:.1f}%</span>
                        </div>
                        """, unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)
            
            if analysis_data.get('recommendations') and analysis_data['recommendations'].get('strategies'):
                with st.expander("📚 Learning Recommendations", expanded=True):
                    st.markdown("""
                    <div class="recommendation-card"><div class="analysis-header"><i>📚</i> Personalized study strategies:</div>
                    """, unsafe_allow_html=True)
                    for strategy in analysis_data['recommendations']['strategies'][:3]:
                        st.markdown(f"<div class='strategy-item'>{strategy}</div>", unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)
            
            if analysis_data.get('learning_patterns'):
                with st.expander("📈 Learning Patterns", expanded=False):
                    st.markdown(f"""
                    <div class="analysis-header"><i>📈</i> Your learning patterns:</div>
                    <p>Time vs Score Correlation: <strong>{analysis_data['learning_patterns']['time_score_correlation']:.2f}</strong></p>
                    <p>Most effective study duration: <strong>{analysis_data['learning_patterns']['optimal_duration']} minutes</strong></p>
                    <p>Best time for studying: <strong>{analysis_data['learning_patterns']['optimal_time']}</strong></p>
                    """, unsafe_allow_html=True)
                    
                    if analysis_data['learning_patterns'].get('suggested_adjustments'):
                        st.markdown("""
                        <div class="analysis-header"><i>🔄</i> Suggested adjustments:</div>
                        """, unsafe_allow_html=True)
                        for adjustment in analysis_data['learning_patterns']['suggested_adjustments'][:2]:
                            st.markdown(f"<div class='strategy-item'>{adjustment}</div>", unsafe_allow_html=True)
        else:
            st.info("Complete some assessments to get personalized learning analysis." 
                    if not analytics.get('assessments') 
                    else "We found your assessment results but couldn't generate recommendations. Please try again.")

# Learning Path Page(xxlj)
def show_learning_path():
    # Initialize the backend timing
    init_backend_timer()
    # Update and get the total duration
    total_minutes = get_total_study_time_from_backend()
    total_hours = total_minutes / 60
    
    user = st.session_state.user
    path = st.session_state.active_path
    paths = learning_engine.get_learning_paths(user['id'])
    if not path:
        st.error("No active path selected")
        if st.button("Back to dashboard"):
            st.session_state.current_view = 'dashboard'
            st.rerun()
        return
    
    # Display the real-time total duration
    st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m", delta="There may be delays.")
    path_id = path['id']
    
    try:
        content = json.loads(path['content'])
        topics = content.get('topics', [])
        
        
        new_progress = calculate_path_progress(path_id, topics, path)
        if abs(new_progress - path['progress']) > 0.01:
            learning_engine.update_learning_progress(path_id, user['id'], new_progress)
            st.session_state.active_path['progress'] = new_progress
        
        resource_counts = st.session_state.path_resource_counts.get(path_id, {'total': 0, 'viewed': 0})
        topic_counts = st.session_state.path_topic_counts.get(path_id, {'total': 0, 'completed': 0})
        
        st.subheader(f"Learning Path: {path['subject']}")
        
        if path.get('ai_generated', False):
            st.info("✨ This learning path was generated by AI to match your specific learning needs")
        
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"Difficulty: {path['difficulty_level']}")
            st.write(f"Target completion: {path['target_completion_date']}")
        with col2:
            st.progress(path['progress'])
            st.write(f"Progress: {path['progress']*100:.1f}%")
        
        # Progress Details
        with st.expander("Progress Details"):
            col_a, col_b = st.columns(2)
            with col_a:
                st.write(f"Resources viewed: {resource_counts['viewed']}/{resource_counts['total']}")
                if resource_counts['total'] > 0: st.progress(resource_counts['viewed'] / resource_counts['total'])
            with col_b:
                st.write(f"Topics completed: {topic_counts['completed']}/{topic_counts['total']}")
                if topic_counts['total'] > 0: st.progress(topic_counts['completed'] / topic_counts['total'])
            
            if topics and path_id in st.session_state.topic_progress:
                st.markdown("### Topic Progress Breakdown")
                topic_progress_list = st.session_state.topic_progress[path_id]
                for i, (topic, progress) in enumerate(zip(topics, topic_progress_list)):
                    topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                    score = st.session_state.topic_assessment_scores.get(topic_key, 0)
                    status = "✓ Completed" if progress >= 0.99 else "In Progress"
                    score_text = f"Assessment: {score}%" if score > 0 else "Assessment not completed"
                    st.write(f"Topic {i+1}: {topic['name']} - {progress*100:.1f}% - {status}")
                    st.write(f"  {score_text}")
                    st.progress(progress)
            
            st.write("*Overall progress is based on completed topics (requires assessment score ≥80%)*")
        
        # Upload Materials(scwj)
        if path_id in st.session_state.uploaded_materials:
            st.subheader("Uploaded Study Materials")
            for mat_idx, material in enumerate(st.session_state.uploaded_materials[path_id]):
                st.markdown("---")
                st.write(f"**📄 {material['name']}**")
                st.write(f"Type: {material['type']}")
                st.write(f"Description: {material['description']}")
                
                if material.get('text_preview') and st.checkbox("View Content Preview", key=f"preview_{material['id']}_{mat_idx}"):
                    st.text_area("Content Preview", material['text_preview'], height=200, disabled=True,
                                key=f"preview_area_{material['id']}_{mat_idx}")
                
                if material.get('file_data'):
                    st.download_button("Download File", material['file_data'], file_name=material['name'],
                                      mime=material['mime_type'], key=f"mat_download_{material['id']}_{mat_idx}")
        
        # Hashtag Page
        if topics:
            topic_tabs = st.tabs([f"Topic {i+1}: {t['name']}" for i, t in enumerate(topics)])
            for i, (topic, tab) in enumerate(zip(topics, topic_tabs)):
                with tab:
                    st.subheader(topic['name'])
                    st.write(topic['description'])
                    st.write(f"Duration: {topic.get('duration_days', 3)} days")
                    
                    if path_id in st.session_state.topic_progress and i < len(st.session_state.topic_progress[path_id]):
                        topic_progress = st.session_state.topic_progress[path_id][i]
                        st.progress(topic_progress)
                        status = "Completed" if topic_progress >= 0.99 else "In Progress"
                        st.write(f"Topic Progress: {topic_progress*100:.1f}% ({status})")
                    
                    topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                    state_content = learning_engine.get_assessments_by_topic(user['id'], path['subject'], topic['name'])
                    if state_content and state_content.get('content'):
                        scores = json.loads(state_content['content'])['scores']
                        valid_scores = [s for s in scores if isinstance(s, (int, float))]
                        score = (sum(valid_scores)/len(valid_scores)*100) if valid_scores else 0.0
                        if score >= 80:
                            st.success(f"Assessment passed! Score: {score}%")
                        elif score > 0:
                            st.warning(f"Assessment score: {score}% (Need 80% to complete this topic)")
                        else:
                            st.info("No assessment completed yet")
                    
                    # Resource List(zygl yck)
                    st.subheader("Resources")
                    for j, res in enumerate(topic.get('resources', [])):
                        st.markdown("---")
                        st.write(f"**{res.get('title', 'Resource')}**")
                        st.write(f"Type: {res.get('type', 'Unknown')}")
                        st.write(res.get('description', 'No description available'))
                        
                        if res.get('platform') == "Uploaded" and res.get('url', '').startswith("uploaded_"):
                            material_id = res['url'].split("_")[1]
                            materials = st.session_state.uploaded_materials.get(path_id, [])
                            material = next((m for m in materials if m['id'] == material_id), None)
                            
                            if material and material.get('file_data'):
                                st.download_button("Download File", material['file_data'], file_name=material['name'],
                                                  mime=material['mime_type'], key=f"res_download_{material_id}_{i}_{j}")
                        elif res.get('url') and res['url'] not in ['#', '']:
                            st.link_button("Open Resource", res['url'])
                        
                        # Mark viewed
                        resource_key = f"{path_id}_{i}_{j}"
                        is_viewed = learning_engine.check_resource_viewed(user['id'], path_id, topic['name'], res['title'])
                        
                        if is_viewed:
                            st.markdown("""<style>.viewed-btn {background-color: #cccccc !important; color: #666666 !important; pointer-events: none;}</style>""", unsafe_allow_html=True)
                            st.button("✅ Marked as Viewed", key=f"viewed_{resource_key}", disabled=True, help="You've already viewed this resource")
                        else:
                            if st.button("Mark as viewed", key=f"view_{resource_key}"):
                                learning_engine.update_viewed_resource(user['id'], path_id, topic['name'], res["title"])
                                st.success("Resource marked as viewed!")
                                new_progress = calculate_path_progress(path_id, topics, path)
                                learning_engine.update_learning_progress(path_id, user['id'], new_progress)
                                st.session_state.active_path['progress'] = new_progress
                                st.rerun()
                    
                    # Evaluation Module
                    st.markdown("---")
                    st.subheader("Topic Mastery Assessment")
                    
                    topic_key = f"{path_id}_topic_{i}"
                    if topic_key not in st.session_state.assessment_generating:
                        st.session_state.assessment_generating[topic_key] = False
                    if topic_key not in st.session_state.show_assessment:
                        st.session_state.show_assessment[topic_key] = False
                    
                    st.write("You must score 80% or higher to complete this topic:")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1: use_saved = st.button("The Last Assessment", key=f"use_saved_{i}")
                    with col2: generate_new = st.button("Re-Assessment", key=f"generate_new_{i}")
                    with col3: save_current = st.button("Save Assessment", key=f"save_current_{i}")
                    with col4: delete_saved = st.button("Delete Saved", key=f"delete_saved_{i}")
                    # scxtm
                    if generate_new:
                        if st.session_state.assessment_generating[topic_key]:
                            st.info("The system has crashed. Please do not click multiple times. It is recommended to log in again")
                        else:
                            st.session_state.assessment_generating[topic_key] = True
                            
                            if not st.session_state.api_key:
                                st.error("DeepSeek API Key is required. Please set it in the dashboard.")
                                st.session_state.assessment_generating[topic_key] = False
                            else:
                                with st.spinner("Generate a new assessment..."):
                                    full_topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                                    
                                    # 1. Clear this topic from the Session
                                    for key in list(st.session_state.keys()):
                                        if full_topic_key in key:
                                            del st.session_state[key]
                                    
                                    # 2. Delete this topic from the database
                                    learning_engine.data_manager.execute_query(
                                        "DELETE FROM assessments WHERE user_id = %s AND topic_name = %s",
                                        (user['id'], topic['name'])
                                    )
                                    
                                    # 3. Force refresh the path data
                                    st.session_state.active_path = None  # Clear the cache of the old path
                                    st.session_state.active_path = learning_engine.get_learning_path(path_id, user['id'])  # 重新加载
                                    
                                    # "Generate logic
                                    if full_topic_key in st.session_state.assessment_state:
                                        del st.session_state.assessment_state[full_topic_key]
                                    
                                    questions_result = assessment_manager.generate_practice_exercises(
                                        path['subject'], topic['name'], path['difficulty_level'],
                                        st.session_state.ai_agent, num_exercises=10
                                    )
                                    
                                    if not questions_result or questions_result["status"] != "success" or len(questions_result["exercises"]) == 0:
                                        st.error("Failed to generate assessment questions. Please try again.")
                                        st.session_state.assessment_generating[topic_key] = False
                                    else:
                                        questions = questions_result["exercises"]
                                        topic['questions'] = questions
                                        # After the assessment is generated, the path data cached at the front end is updated synchronously
                                        learning_engine.add_topic_questions(path['id'], paths, topic['name'], questions, user['id'])
                                        # Re-obtain the latest path data and update it to the session state
                                        st.session_state.active_path = learning_engine.get_learning_path(path_id, user['id'])
                                        # Re-parse the path content (make sure to include the newly generated evaluation question)
                                        content = json.loads(st.session_state.active_path['content'])
                                        topics = content.get('topics', [])  # Refresh the topics list
                                        st.session_state.show_assessment[topic_key] = True

                                        st.session_state.assessment_generating[topic_key] = False
                                        st.success(f"The new assessment has been generated successfully! {topic['name']} ")
                                        st.rerun()

                    if save_current:
                        full_topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                        if 'questions' in topic and isinstance(topic['questions'], list) and len(topic['questions']) > 0:
                            st.session_state.saved_assessments[full_topic_key] = {
                                'questions': topic['questions'], 'saved_at': datetime.now().strftime("%Y-%m-%d %H:%M")
                            }
                            st.success("Assessment saved successfully!")
                        else:
                            st.warning("No assessment questions to save. Please generate a new assessment first.")

                    if delete_saved:
                        full_topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                        if full_topic_key in st.session_state.saved_assessments:
                            del st.session_state.saved_assessments[full_topic_key]
                        if full_topic_key in st.session_state.assessment_state:
                            del st.session_state.assessment_state[full_topic_key]
                        if 'questions' in topic: del topic['questions']
                        learning_engine.add_topic_questions(path['id'], paths, topic['name'], "", user['id'])
                        st.session_state.show_assessment[topic_key] = False
                        if full_topic_key in st.session_state.topic_assessment_scores:
                            del st.session_state.topic_assessment_scores[full_topic_key]
                        st.session_state.assessment_results = [
                            res for res in st.session_state.assessment_results 
                            if res.get('topic_key') != full_topic_key
                        ]
                        new_progress = calculate_path_progress(path_id, topics, path)
                        learning_engine.update_learning_progress(path_id, user['id'], new_progress)
                        st.session_state.active_path['progress'] = new_progress
                        st.success("Saved assessment and all questions have been deleted.")
                        st.rerun()
                    # else:
                        # st.warning("No saved assessment to delete.")

                    if use_saved:
                        questions = topic.get('questions', [])
                        if questions and isinstance(questions, list) and len(questions) > 0:
                            st.session_state.show_assessment[topic_key] = True
                            st.rerun()
                        else:
                            st.warning("No saved assessment available for this topic.")
                            st.session_state.show_assessment[topic_key] = False
                    else:
                        questions = topic.get('questions', []) if st.session_state.show_assessment.get(topic_key, False) else []

                    if questions and isinstance(questions, list) and len(questions) > 0 and st.session_state.show_assessment.get(topic_key, False):
                        st.markdown(f"### {topic['name']} Assessment (Total {len(questions)} Questions)")
                        st.info("You need to score at least 80% to complete this topic.")
                        
                        full_topic_key = f"{path_id}_topic_{i}_{topic['name']}"
                        if full_topic_key not in st.session_state.assessment_state:
                            st.session_state.assessment_state[full_topic_key] = {
                                'user_answers': ["" for _ in questions], 'scores': [], 'feedback': [],
                                'submitted': False, 'questions': questions
                            }
                        
                        current_state = st.session_state.assessment_state[full_topic_key]
                        state_content = learning_engine.get_assessments_by_topic(user['id'], path['subject'], topic['name'])
                        if state_content and not current_state['submitted']:
                            saved_state = json.loads(state_content.get('content', '{}'))
                            if saved_state.get('user_answers') and len(saved_state['user_answers']) == len(questions):
                                current_state['user_answers'] = saved_state['user_answers']
                                current_state['questions'] = questions
                                current_state['submitted'] = True
                                current_state['scores'] = json.loads(state_content['content'])['scores']
                                current_state['feedback'] = json.loads(state_content['content'])['feedback']
                                valid_scores = [s for s in current_state['scores'] if isinstance(s, (int, float))]
                                current_state['total_score'] = (sum(valid_scores)/len(valid_scores)*100) if valid_scores else 0.0
                        # (pghd)
                        with st.form(f"assessment_form_{full_topic_key}", clear_on_submit=False):
                            for idx, q in enumerate(questions):
                                with st.container():
                                    st.markdown(f"""
                                    <div class="question-container">
                                        <div class="question-header">
                                            <div class="question-number">Question {idx+1}</div>
                                            <div class="question-difficulty">{q.get('difficulty', 'Medium')}</div>
                                        </div>
                                        <div class="question-content">{q.get('question', 'No question content')}</div>
                                    """, unsafe_allow_html=True)
                                    
                                    current_answer = current_state['user_answers'][idx]
                                    if q.get('type') == 'multiple_choice' and q.get('options'):
                                        option_index = q['options'].index(current_answer) if current_answer in q['options'] else None
                                        user_answer = st.radio(
                                            "Select your answer:", q['options'], index=option_index,
                                            key=f"q_radio_{full_topic_key}_{idx}"
                                        )
                                    else:
                                        user_answer = st.text_area(
                                            "Your answer:", value=current_answer, height=100,
                                            key=f"q_textarea_{full_topic_key}_{idx}"
                                        )
                                    
                                    current_state['user_answers'][idx] = user_answer if user_answer is not None else ""
                                    st.markdown("</div>", unsafe_allow_html=True)
                            
                            submit_assessment = st.form_submit_button("Submit Assessment")
                            if submit_assessment:
                                if not st.session_state.api_key:
                                    st.error("DeepSeek API Key is required. Please set it in the dashboard.")
                                    return
                                all_answered = all(ans.strip() != "" for ans in current_state['user_answers'])
                                if not all_answered:
                                    st.warning("Please answer all questions before submitting!")
                                else:
                                    with st.spinner("Evaluating your answers..."):
                                        current_state['scores'] = []
                                        current_state['feedback'] = []
                                        for idx, (q, ans) in enumerate(zip(questions, current_state['user_answers'])):
                                            evaluation = assessment_manager.evaluate_answer(
                                                path['subject'], topic['name'], q['question'],
                                                ans, path['difficulty_level'], st.session_state.ai_agent
                                            )
                                            score = evaluation['data']['score'] if (evaluation and evaluation["status"] == "success") else 0.0
                                            feedback = evaluation['data']['feedback'] if (evaluation and evaluation["status"] == "success") else "No feedback available"
                                            current_state['scores'].append(score)
                                            current_state['feedback'].append(feedback)
                                        
                                        valid_scores = [s for s in current_state['scores'] if isinstance(s, (int, float))]
                                        current_state['total_score'] = (sum(valid_scores)/len(valid_scores)*100) if valid_scores else 0.0
                                        current_state['submitted'] = True
                                        
                                        # Explicitly update the top-level session_state key
                                        st.session_state.assessment_state[full_topic_key] = current_state
                                        
                                        st.session_state.topic_assessment_scores[full_topic_key] = current_state['total_score']
                                        
                                        assessment_result = {
                                            'user_id': user['id'], 'topic_key': full_topic_key,
                                            'subject': path['subject'], 'topic': topic['name'],
                                            'score': current_state['total_score'],
                                            'date': datetime.now().strftime("%Y-%m-%d %H:%M"),
                                            'questions_count': len(questions)
                                        }
                                        
                                        found_index = None
                                        for idx, res in enumerate(st.session_state.assessment_results):
                                            if res.get('topic_key') == full_topic_key:
                                                found_index = idx
                                                break
                                        
                                        if found_index is not None:
                                            st.session_state.assessment_results[found_index] = assessment_result
                                        else:
                                            st.session_state.assessment_results.append(assessment_result)
                                        
                                        new_progress = calculate_path_progress(path_id, topics, path)
                                        learning_engine.update_learning_progress(path_id, user['id'], new_progress)
                                        st.session_state.active_path['progress'] = new_progress
                                        
                                        st.rerun()

                        if current_state['submitted']:
                            st.markdown("---")
                            st.markdown(f"""
                            <div class="assessment-container">
                                <div class="assessment-header">
                                    <div class="assessment-title">Assessment Results</div>
                                    <div class="assessment-score">{current_state['total_score']:.1f}%</div>
                                </div>
                            """, unsafe_allow_html=True)
                            
                            if current_state['total_score'] >= 80:
                                st.success("Congratulations! You passed this assessment and completed the topic!")
                            else:
                                st.warning(f"You need at least 80% to complete this topic. You scored {current_state['total_score']:.1f}%.")
                            
                            for idx, (q, ans, score, feedback) in enumerate(zip(
                                questions, current_state['user_answers'], current_state['scores'], current_state['feedback']
                            )):
                                with st.expander(f"Question {idx+1}: {q.get('question', 'No question')}", expanded=False):
                                    display_ans = ans.strip() if ans.strip() != "" else "Not filled in"
                                    st.write(f"**Your answer:** {display_ans}")
                                    
                                    if q.get('type') == 'multiple_choice' and q.get('options') and q.get('correct_option') is not None:
                                        correct_idx = q['correct_option']
                                        if 0 <= correct_idx < len(q['options']):
                                            st.write(f"**Correct answer:** {q['options'][correct_idx]}")
                                    
                                    st.write(f"**Score:** {score*100:.1f}%")
                                    st.write(f"**Feedback:** {feedback}")
                            
                            st.markdown("</div>", unsafe_allow_html=True)
                            learning_engine.insert_assessment_from_state(user['id'], path.get('subject', 'Unknown'), topic['name'], current_state)
                            st.success("Assessment submitted successfully!")
                    # elif not generate_new or not use_saved:
                        # st.info("Please generate a new assessment or use a saved one to begin.")
        
        # Milestones
        milestones = content.get('milestones', [])
        if milestones:
            st.subheader("Key Milestones")
            for milestone in milestones:
                with st.expander(f"Day {milestone['expected_completion_day']}: {milestone['name']}"):
                    st.write(f"**Criteria:** {milestone['assessment_criteria']}")
        
        # Path Completion judgment(ljwc)
        if path['progress'] >= 0.99:
            path_ids = [p['id'] for p in st.session_state.completed_paths]
            if path['id'] not in path_ids:
                st.session_state.completed_paths.append(path)
                st.balloons()
                st.success("🎉 Congratulations! You've completed this learning path!")
                
                if path['id'] not in st.session_state.achievement_path_ids:
                    achievement = {
                        "title": f"Completed {path['subject']}",
                        "description": f"Successfully finished the {path['subject']} learning path",
                        "date": datetime.now().strftime("%Y-%m-%d"),
                        "score": round(path['progress'] * 100)
                    }
                    st.session_state.achievements.append(achievement)
                    st.session_state.achievement_path_ids.add(path['id'])
        
        # Achievement Display
        st.subheader("My Achievements")
        if not st.session_state.achievements:
            st.info("No achievements yet. Complete learning activities to earn achievements!")
        else:
            for achievement in st.session_state.achievements:
                with st.expander(f"{achievement['title']} - {achievement['date']}", expanded=False):
                    st.write(achievement['description'])
                    if achievement.get('score'):
                        st.write(f"Score: {achievement['score']}%")
        
        # Return to the dashboard
        if st.button("Back to dashboard"):
            st.session_state.current_view = 'dashboard'
            st.rerun()
        
    except Exception as e:
        st.error(f"Error: {str(e)}")
        if st.button("Back"):
            st.session_state.current_view = 'dashboard'
            st.rerun()

# Create a path page
def show_create_path():
    # Initialize the backend timing
    init_backend_timer()
    # Update and get the total duration
    total_minutes = get_total_study_time_from_backend()
    total_hours = total_minutes / 60
    
    # Display the real-time total duration of 2
    st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m", delta="There may be delays.")

    user = st.session_state.user
    if not user:
        st.session_state.current_view = 'login'
        st.rerun()
        
    st.subheader("Create New Learning Path")
    
    # Create a path form
    with st.form("create_path_form"):
        col1, col2 = st.columns(2, gap="small")
        with col1:
            subject = st.text_input("Subject")
            difficulty = st.select_slider("Difficulty", ["Beginner", "Intermediate", "Advanced"])
        with col2:
            target_days = st.number_input("Target days", 7, 90, 30)
        
        user_profile = user_manager.get_user_profile(user['id'])
        interests = st.text_area("Interests", user_profile['interests'], height=150)
        learning_style = st.selectbox(
            "Learning style", ["Visual", "Auditory", "Reading/Writing", "Kinesthetic"],
            index=["Visual", "Auditory", "Reading/Writing", "Kinesthetic"].index(user_profile['learning_style'])
        )
        
        st.subheader("Upload Study Materials (Optional)")
        uploaded_files = st.file_uploader(
            "Upload PDF, Word, or Text files", accept_multiple_files=True,
            type=["pdf", "docx", "txt"]
        )
        
        materials = []
        if uploaded_files:
            for file_idx, file in enumerate(uploaded_files):
                material_id = str(uuid.uuid4())[:8]
                file_ext = file.name.split(".")[-1].lower()
                file_type = {"pdf": "PDF Document", "docx": "Word Document", "txt": "Text File"}.get(file_ext, "Document")
                text_preview = extract_text_from_file(file)
                file_data = file.getvalue()
                mime_type = {"pdf": "application/pdf", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                            "txt": "text/plain"}.get(file_ext, "application/octet-stream")
                
                desc = st.text_input(
                    f"Description for {file.name}", f"Uploaded {file_type} for {subject}",
                    key=f"desc_{material_id}_{file_idx}"
                )
                actual_text = text_preview.get('text', '')
                text_preview_processed = actual_text[:500] + ("..." if len(actual_text) > 500 else "")
                materials.append({
                    "id": material_id, "name": file.name, "type": file_type, "description": desc,
                    "mime_type": mime_type, "text_preview": text_preview_processed,
                    "file_data": file_data
                })
        
        ai_option = st.checkbox("Use AI to generate personalized learning path", value=True)
        submit = st.form_submit_button("Generate Path")
        
        if submit:
            if not subject:
                st.warning("Please enter a subject")
            else:
                with st.spinner("Generating your personalized learning path..."):
                    if not st.session_state.api_key:
                        st.error("DeepSeek API Key is required. Please set it in the dashboard.")
                        return
                    path_id, default = learning_engine.create_learning_path(
                        user['id'], subject, difficulty, target_days, st.session_state.ai_agent
                    )
                    if not default:
                        st.error("DeepSeek error, please check it!")
                        return
                    
                    if path_id:
                        if materials:
                            st.session_state.uploaded_materials[path_id] = materials
                        
                        st.session_state.active_path = learning_engine.get_learning_path(path_id, user['id'])
                        st.session_state.selected_path_id = path_id
                        st.session_state.ai_generated_path = ai_option
                        st.session_state.show_assessment = {}
                        # if ai_option:
                        #     st.warning("Note: AI generation is currently unavailable. The path uses default content.")
                        st.success("Path created successfully!")
                        st.session_state.current_view = 'learning_path'
                        st.rerun()
                    else:
                        st.error("Failed to create path. Try again.")

# Analysis Page
def show_analytics():
    try:  # Add global exception catching
        # Initialize the backend timing
        init_backend_timer()
        # Update and get the total duration
        total_minutes = get_total_study_time_from_backend()
        total_hours = total_minutes / 60
        
        # Display the real-time total duration of 3
        st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m", delta="There may be delays.")

        user = st.session_state.user
        if not user:
            st.session_state.current_view = 'login'
            st.rerun()
            
        st.subheader("Your Learning Analytics")
        
        # Obtain analysis data from the back end
        analytics = learning_engine.get_learning_analytics(
            user['id'], st.session_state.completed_paths
        )
        # Make sure analytics is of dictionary type
        if not isinstance(analytics, dict):
            analytics = {}
        # Supplement the total duration to analytics
        analytics['total_study_time'] = total_minutes
        
        # Process path data
        paths = analytics.get('paths', [])
        if paths and isinstance(paths, list):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Paths", len(paths))
            with col2:
                try:
                    avg_prog = sum(p.get('progress', 0) for p in paths) / len(paths)
                    st.metric("Avg Progress", f"{avg_prog*100:.1f}%")
                except (ZeroDivisionError, TypeError):
                    st.metric("Avg Progress", "N/A")
           
        # Process activity data
        activities = analytics.get('activities', [])
        if activities and isinstance(activities, list) and 'col3' in locals():
            try:
                activity_df = pd.DataFrame(activities)
                if 'topic_name' in activity_df.columns and 'total_minutes' in activity_df.columns:
                    topic_totals = activity_df.groupby('topic_name')['total_minutes'].sum()
                    if not topic_totals.empty:
                        top_topic = topic_totals.idxmax()
                        top_hours = topic_totals.max() / 60
                        with col3:
                            st.subheader("Most Studied")
                            st.write(f"{top_topic}")
                            st.write(f"≈ {top_hours:.1f} hours")
            except Exception as e:
                st.warning(f"Error processing activity data: {str(e)}")
          
        # Progress Chart
        if paths:
            st.subheader("Progress Over Time")
            try:
                chart_result = LearningAnalytics.generate_progress_chart(paths)
                if isinstance(chart_result, dict) and chart_result.get("status") == "success" and "figure" in chart_result:
                    st.plotly_chart(chart_result["figure"], use_container_width=True)
                else:
                    st.info("Failed to generate progress chart (invalid data)")
            except Exception as e:
                st.warning(f"Error generating progress chart: {str(e)}")

        # Evaluate the radar chart
        assessments = analytics.get('assessments', [])
        if assessments and isinstance(assessments, list):
            st.subheader("Assessment Performance")
            try:
                radar_result = LearningAnalytics.generate_assessment_radar(assessments)
                if isinstance(radar_result, dict) and radar_result.get("status") == "success" and "figure" in radar_result:
                    st.plotly_chart(radar_result["figure"], use_container_width=True)
                else:
                    st.info("Failed to generate assessment radar (invalid data)")
            except Exception as e:
                st.warning(f"Error generating radar chart: {str(e)}")
        else:
            st.info("No enough assessment data to generate radar chart")
          
        # Learning Mode Analysis
        if activities or assessments:  # Any data will be displayed
            st.subheader("Learning Patterns")
            try:
                has_enough_data = (
                    len(activities) >= 3 and  # At least 3 activity records
                    len(assessments) >= 2    # At least two assessment records
                )
                
                if has_enough_data:
                    patterns = LearningAnalytics.identify_learning_patterns(activities, assessments)
                else:
                    patterns = {}  # Manually build an empty result when data is insufficient
                
                with st.expander("Details", expanded=True):
                    # 1. Handle topic scores
                    if patterns.get('top_topics'):
                        st.write("**Strongest topics (by score):**")
                        for topic in patterns['top_topics'][:3]:
                            st.write(f"- {topic.get('topic', 'Unknown')} (Avg: {topic.get('avg_score', 0)*100:.1f}%)")
                    else:
                        st.info("No enough assessment data to determine strong topics")
                        
                    # 2. Processing time-fraction Correlation (Level 1: Within expander)
                    corr = patterns.get('time_score_correlation', None)
                    if corr is not None and has_enough_data:
                        st.write(f"**Study time vs score correlation:** {corr:.2f}")
                        # Add Relevant explanations
                        if abs(corr) < 0.3:
                            st.caption("Weak or no correlation between study time and scores")
                        elif 0.3 <= abs(corr) < 0.7:
                            st.caption("Moderate correlation between study time and scores")
                        else:
                            st.caption("Strong correlation between study time and scores")
                    else:
                        st.info("Not enough data to calculate correlation")
                        
            except Exception as e:
                st.warning(f"Error analyzing learning patterns: {str(e)}")
                
        # Download Report
        if st.button("Download Detailed Study Report", use_container_width=True):
            try:
                with st.spinner("Generating PDF report..."):
                    # Make sure the PDF generator can obtain the correct data
                    if isinstance(analytics, dict):
                        report_result = PDFGenerator.generate_study_report(
                            user['id'], user['username'], analytics
                        )
                        if report_result["status"] == "success":
                            st.download_button(
                                label="Save PDF Report",
                                data=report_result["buffer"],
                                file_name=f"study_analytics_{datetime.now():%Y%m%d_%H%M%S}.pdf",
                                mime="application/pdf",
                                key="download_analytics_report"
                            )
                        else:
                            st.error(f"Failed to generate report: {report_result['message']}")
                    else:
                        st.error("Invalid analytics data for report generation")
            except Exception as e:
                st.error(f"Failed to generate report: {str(e)}")
    
    except Exception as e:  # Catch all unhandled exceptions
        st.error(f"Analytics page error: {str(e)}")
        # Stay on the current page and do not jump back to the Dashboard
        st.session_state.current_view = 'analytics'

# Plan Page
def show_planner():
    # Helper Function: Generate a study plan in Word format
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    def generate_study_plan_word(plan_data):
        # 1. Create an empty Word document
        doc = Document()

        # 2. Set the document title
        title = doc.add_heading(f"Study Plan：{plan_data['basic_info']['subject']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3. Add basic plan information
        basic_info = plan_data["basic_info"]
        info_para = doc.add_paragraph()
        info_para.add_run(f"study time：{basic_info['daily_hours']}\n").bold = True
        info_para.add_run(f"Planned deadline：{basic_info['deadline']}\n").bold = True
        info_para.add_run(f"Generation time：{basic_info['generated_time']}\n").bold = True
        doc.add_paragraph()  # Blank line separation

        # 4. Daily study schedule
        doc.add_heading("I. Daily Study Schedule", 1)
        daily_schedule = plan_data["daily_schedule"]

        for day in daily_schedule:
            # Daily title (such as "Monday")
            doc.add_heading(f"{day['day']}", 2)

            # Create a 3-column table: study type, topic, duration/focus
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"  # The table has borders.
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Learning type"
            hdr_cells[1].text = "Learning topics"
            hdr_cells[2].text = "Duration (minutes)/Key areas"

            # Fill in the learning block data of the day
            for block in day["study_blocks"]:
                # Automatically determine the type of learning (knowledge learning/practice consolidation)
                block_type = "Knowledge learning" if "Concept" in block["focus_area"].lower() else "Practice and consolidation"
                row_cells = table.add_row().cells
                row_cells[0].text = block_type
                row_cells[1].text = block["topic"]
                row_cells[2].text = f"{block['duration_minutes']}minutes | {block['focus_area']}"

            doc.add_paragraph()  # Empty walk after each day's schedule

        # 5. Learning Efficiency Tips
        doc.add_heading("Ii. Tips on Learning Efficiency", 1)
        tips = plan_data["productivity_tips"]
        if tips:
            for tip in tips:
                doc.add_paragraph(tip, style="List Bullet")
        else:
            doc.add_paragraph("There is no efficiency prompt for now. You can add it manually", style="List Bullet")

        # 6. Save the document to memory
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)  # The pointer returns to the beginning.
        return doc_bytes

    # Initialize the backend timing + obtain the real-time total duration
    init_backend_timer()
    total_minutes = get_total_study_time_from_backend()
    total_hours = total_minutes / 60
    
    user = st.session_state.user
    if not user:
        st.session_state.current_view = 'login'
        st.rerun()
    
    # Display the real-time total duration
    st.metric("Current path learning time", f" {total_minutes // 60}h {total_minutes % 60}m", delta="There may be delays.")
    st.subheader("📅 Study Planner")
    
    # Obtain the user's existing learning path
    paths = learning_engine.get_learning_paths(user['id'])
    if not paths:
        st.warning("No learning paths found! Create a path first to generate a study plan.")
        if st.button("Go to Create Path", use_container_width=True):
            st.session_state.current_view = 'create_path'
            st.rerun()
        return
    
    # Path selection drop-down box
    path_options = {str(path['id']): path['subject'] for path in paths}
    selected_path_id = st.selectbox(
        "Select a learning path to focus on",
        options=list(path_options.keys()),
        format_func=lambda x: path_options[x]
    )
    selected_path = next(p for p in paths if str(p['id']) == selected_path_id)
    # The plan for storing the currently selected path (for subsequent saving/deleting)
    current_plan = None
    
    # Plan parameter Settings
    col1, col2 = st.columns(2, gap="medium")
    with col1:
        # Daily study duration (hours to minutes)
        daily_hours = st.slider(
            "Daily study time",
            min_value=0.5, max_value=8.0, value=2.0, step=0.5,
            key="daily_hours_slider"
        )
        daily_minutes = int(daily_hours * 60)
        
        # Deadline
        deadline = st.date_input(
            "Plan deadline",
            value=datetime.now() + timedelta(days=7),
            min_value=datetime.now(),
            key="plan_deadline"
        )
    
    with col2:
        # Key areas of focus
        focus_areas = st.text_input(
            "Focus areas (e.g., basics, exercises)",
            value=selected_path['subject'] + " basics",
            key="focus_areas_input"
        )
        
        # Plan generation method
        plan_type = st.radio(
            "Plan type",
            options=["Balanced (mix of learning & practice)", "Learning-focused", "Practice-focused"],
            key="plan_type_radio"
        )
    
    # Generate the plan + explicitly save the logic
    if st.button("Generate Personalized Study Plan", use_container_width=True):
        with st.spinner("Generating your study plan..."):
            # Extract the topic from the path
            try:
                path_content = json.loads(selected_path['content'])
                topics = [t['name'] for t in path_content.get('topics', [])]
            except json.JSONDecodeError:
                topics = [f"Topic {i+1}" for i in range(3)]  # Use the default theme when parsing fails
            
            # # U se the default theme when parsing fails()
            ai_plan_result = st.session_state.ai_agent.generate_study_schedule(
                deadline=deadline,
                hours_per_day=daily_hours,
                topics=topics,
                subject=selected_path['subject'],
                focus=focus_areas
            )
            # Extract the planned data returned by AI
            ai_plan = ai_plan_result.get("data", {}) if ai_plan_result.get("status") == "success" else {}
            
            # After generating the plan, only display and save the original plan
            if ai_plan.get('daily_schedule'):
                # Automatically save the plan to the database
                save_result = learning_engine.insert_plan_from_json(
                    user_id=user['id'],
                    path_id=int(selected_path_id),
                    study_schedules=ai_plan
                )
                current_plan = ai_plan
                st.session_state[f"current_plan_{selected_path_id}"] = ai_plan

                # Display the generated original plan
                st.subheader(f"📋 {selected_path['subject']} Study Plan (Deadline: {deadline.strftime('%Y-%m-%d')})")
                st.write(f"Total duration: {daily_hours} hours/day for {(deadline - datetime.now().date()).days} days")
            
                # Display the original plan on a daily basis
                for day in ai_plan['daily_schedule']:
                    with st.expander(f"🗓️ {day['day']}", expanded=False):
                        day_total = sum(block['duration_minutes'] for block in day['study_blocks'])
                        st.write(f"**Current path learning time: {day_total//60}h {day_total%60}m**")
                        for idx, block in enumerate(day['study_blocks']):
                            block_type = "📚 Learning" if "concept" in block['focus_area'].lower() else "✏️ Practice"
                            st.write(f"{idx+1}. {block_type} - {block['subject']}: {block['topic']}")
                            st.write(f"   • Duration: {block['duration_minutes']} hours")
                            st.write(f"   • Focus: {block['focus_area']}")
                            st.write("---")

                # Keep only the "Save Original Plan" button
                st.markdown("#### 💾 Save Plan")
                if st.button("Save Original Plan to Database", key="save_original_plan"):
                    save_result = learning_engine.insert_plan_from_json(
                        user_id=user['id'],
                        path_id=int(selected_path_id),
                        study_schedules=current_plan
                    )
                    if save_result.get("status") == "success":
                        st.success(f"✅ Plan saved successfully! (Plan ID: {save_result['id']})")
                    else:
                        st.error("× Failed to save plan. Please try again.")

                # Retain the download function
                final_plan = {
                    "daily_schedule": ai_plan['daily_schedule'],
                    "productivity_tips": ai_plan.get('productivity_tips', []),
                    "basic_info": {
                        "subject": selected_path["subject"],
                        "daily_hours": f"{daily_hours}hour",
                        "deadline": deadline.strftime("%Y-%m-%d"),
                        "generated_time": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                }
                # "Download Word"
                word_file = generate_study_plan_word(final_plan)
                st.download_button(
                    "📄 Download ",
                    data=word_file,
                    file_name=f"{selected_path['subject']}_study_plan_{datetime.now():%Y%m%d}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word_plan"
                )
            else:
                st.error("There are problems with the current network.")
    else:
        st.success("Note: In case of API disconnection, the default plan will be used.")

    # Delete the function of saved plans
    st.subheader("📜 Saved Study Plans")
    # Create independent containers to isolate dynamic components
    saved_plans_container = st.container()

    with saved_plans_container:  # All dynamic components are placed in this container (indented to level 1 within the function)
        saved_plan = learning_engine.get_plan(user['id'], int(selected_path_id))
        
        if saved_plan:
            try:
                saved_plan_json = json.loads(saved_plan['schedule_json'])
                st.write(f"**Last updated:** {saved_plan['created_at'].strftime('%Y-%m-%d %H:%M')}")
            
                # Display a summary of historical plans
                if saved_plan_json.get('daily_schedule'):
                    daily_schedule = saved_plan_json['daily_schedule']
                    total_days = len(daily_schedule)
                    if total_days > 0:
                        sample_day = daily_schedule[0]
                        duration_text = "One week" if total_days >7 else f"{total_days}day"
                    else:
                        st.write("**Note:** This plan has no scheduled days.")
                else:
                    st.write("**Note:** This plan has no daily schedule data.")
            
                # "View Saved Plan Button (Details will be rendered only after clicking)
                if st.button("View Saved Plan", key="view_saved_plan_btn"):
                    st.subheader(f"Saved Plan for {path_options[selected_path_id]}")
                    # Display the historical plan by day (add a unique identifier to the label)
                    for day_idx, day in enumerate(saved_plan_json['daily_schedule']):
                        # Embed a unique identifier in the label to avoid Streamlit confusion
                        with st.expander(f"🗓️ {day['day']} (Plan {selected_path_id}-{day_idx})", expanded=False):
                            day_total = sum(block['duration_minutes'] for block in day['study_blocks'])
                            st.write(f"**Total: {day_total//60}h {day_total%60}m**")
                            for block_idx, block in enumerate(day['study_blocks']):
                                st.write(f"{block_idx+1}. {block['subject']}: {block['topic']} ({block['duration_minutes']}m)")
            
                # Delete Button (Enhanced Status cleaning)
                st.markdown("---")
                confirm_delete = st.checkbox("I confirm to delete this plan (cannot be undone)", key="confirm_delete")
                if confirm_delete and st.button("🗑️ Delete Saved Plan", key="delete_saved_plan_btn", type="primary"):
                    with st.spinner("Deleting plan..."):
                        # 1. Delete database records
                        delete_result = learning_engine.data_manager.execute_query(
                            """DELETE FROM study_schedules WHERE user_id = %s AND path_id = %s""",
                            (user['id'], int(selected_path_id))
                        )
                        if delete_result is not False:
                            # 2. Thoroughly clean up all relevant session states
                            plan_keys = [k for k in st.session_state if f"current_plan_{selected_path_id}" in k]
                            for k in plan_keys:
                                del st.session_state[k]
                            # 3. Clear all components from the container (Key point!) Force DOM synchronization
                            saved_plans_container.empty()  # Empty the container and delete all child nodes
                            st.success("✅ Saved plan deleted successfully!")
                            # 4. Refresh the page immediately to avoid residual rendering
                            st.rerun()
                        else:
                            st.error("❌ Failed to delete plan. Please try again.")
            
                # Download the historical plan (render only when the plan exists)
                plan_with_info = {
                    "daily_schedule": saved_plan_json.get("daily_schedule", []),
                    "productivity_tips": saved_plan_json.get("productivity_tips", []),
                    "basic_info": {
                        "subject": selected_path["subject"],
                        "daily_hours": f"{daily_hours}hours",
                        "deadline": deadline.strftime("%Y-%m-%d"),
                        "generated_time": saved_plan['created_at'].strftime("%Y-%m-%d %H:%M")
                    }
                }
                word_file = generate_study_plan_word(plan_with_info)
                st.download_button(
                    "Download Saved Plan",
                    data=word_file,
                    file_name=f"{path_options[selected_path_id]}_saved_plan_{saved_plan['created_at'].strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_saved_plan_btn"
                )
            except json.JSONDecodeError:
                st.error("Failed to load saved plan. Please generate a new one.")
                # Delete damaged plan records
                learning_engine.data_manager.execute_query(
                    "DELETE FROM study_schedules WHERE id = %s",
                    (saved_plan['id'],)
                )
        else:
            # Display prompts when there is no plan to avoid rendering empty components
            st.info("No saved study plans yet. Generate a plan to save it.")
        

# Main Application Control
def main():
    # Initialize the session state
    init_session_state()
    # Initialize the AI agent
    init_ai_agent()
    
    # In the main() function, automatic refreshing is only enabled on unplanned pages
    if st.session_state.user and st.session_state.current_view != "planner":
        st_autorefresh(
            interval=300000,  # Change to a 5-minute refresh to reduce conflicts on the plan page
            key="global_auto_refresh",
            debounce=2000  # Increase latency and avoid rapid and continuous refreshing
        )
    
    # Force a jump to the login page when not logged in
    if not st.session_state.user:
        show_login_page()
        return
    
    # Logged in: Display the side navigation bar
    with st.sidebar:
        st.markdown(f"<h2 style='color: {MORANDI_COLORS['accent']};'>📚 Navigation</h2>", unsafe_allow_html=True)
        
        # Navigation menu configuration
        nav_items = [
            ("Dashboard", "dashboard"),
            ("Create New Path", "create_path"),
            ("AI Assistant", "ai_assistant"),
            ("Study Analytics", "analytics"),
            ("Study Planner", "planner")
        ]
        
        # Render navigation button
        for name, view in nav_items:
            # The learning path page requires an active path and is disabled when there is no path
            disabled = (view == "learning_path" and not st.session_state.active_path and not learning_engine.get_learning_paths(st.session_state.user['id']))
            btn_text = f"{name}"
            
            if st.button(btn_text, use_container_width=True, key=f"nav_btn_{view}", disabled=disabled):
                st.session_state.current_view = view
                # Clear the active path when switching to a non-learning path page
                if view != "learning_path":
                    st.session_state.active_path = None
                st.rerun()
        
        # Dividing line
        st.markdown("---")
        
        
        
        # Log out button
        if st.button("Logout", use_container_width=True, key="logout_btn"):
            # When logging out, the database duration is not cleared; only the session status is reset
            st.session_state.user = None
            st.session_state.current_view = 'login'
            st.session_state.timer_initialized = False
            st.success("Logged out successfully!")
            st.rerun()
    
    # Render the corresponding page based on the current view
    current_view = st.session_state.get('current_view', 'dashboard')
    view_mapping = {
        'dashboard': show_dashboard,
        'learning_path': show_learning_path,
        'create_path': show_create_path,
        'ai_assistant': show_ai_assistant,
        'analytics': show_analytics,
        'planner': show_planner
    }
    
    # Execute the function corresponding to the current view
    try:
        view_mapping[current_view]()
    except KeyError:
        # When the view does not exist, it jumps to the dashboard by default
        st.session_state.current_view = 'dashboard'
        st.rerun()

# Program entry
if __name__ == "__main__":
    main()